"""python-docx writer for the press release. We build the document
programmatically rather than via a .j2 template — docx is XML under the
hood and Jinja doesn't round-trip cleanly.

Structure mirrors the reference PDF: 요약 + 01~03 본문 + 04 성과표 + 05 인사이트 + About.
"""
from __future__ import annotations

import datetime as _dt
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

from config import load_settings


OLIVE = RGBColor(0x4E, 0x5A, 0x2D)


def _heading(doc: Document, text: str, *, level: int = 2, size: int = 13):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = OLIVE
    return h


def render_press_docx(context: dict[str, Any], out_path: Path) -> Path:
    s = load_settings()
    campaign = context["campaign"]
    narrative = context["narrative"]
    headline = context["headline"]
    subhead = context.get("subhead", "")
    year = context.get("year", _dt.date.today().year)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)

    eyebrow = doc.add_paragraph()
    r = eyebrow.add_run(f"[보도자료] {s.company_name} Official Case Study")
    r.bold = True
    r.font.color.rgb = OLIVE
    r.font.size = Pt(10)

    title = doc.add_heading(headline, level=1)
    for run in title.runs:
        run.font.size = Pt(20)

    if subhead:
        sub = doc.add_paragraph()
        rs = sub.add_run(subhead)
        rs.italic = True
        rs.font.size = Pt(11)

    if narrative.get("summary"):
        _heading(doc, "요약", size=13)
        doc.add_paragraph(narrative["summary"])

    for key, title_text in [
        ("overview", "01. 캠페인 개요"),
        ("background", "02. 광고 집행 배경"),
        ("strategy", "03. 적용 전략"),
    ]:
        body = (narrative.get(key) or "").strip()
        if not body:
            continue
        _heading(doc, title_text, size=13)
        doc.add_paragraph(body)

    # 04. results table
    metrics = campaign.get("metrics_table") or []
    if metrics:
        _heading(doc, "04. 캠페인 성과", size=13)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "성과 지표"
        hdr[1].text = "성과"
        hdr[2].text = "비고"
        for m in metrics:
            row = table.add_row().cells
            row[0].text = str(m.get("indicator", ""))
            row[1].text = str(m.get("value", ""))
            row[2].text = str(m.get("note", ""))

    insights = narrative.get("insights") or []
    if insights:
        _heading(doc, "05. 인사이트", size=13)
        for item in insights:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(item))

    doc.add_paragraph()
    _heading(doc, f"About {s.company_name}", level=3, size=11)
    doc.add_paragraph(s.company_description)

    footer_p = doc.add_paragraph()
    footer_p.add_run(
        f"Website: {s.company_url}"
        + (f", {s.company_url_secondary}" if s.company_url_secondary else "")
        + f"  |  Contact Us: {s.press_contact_email}"
    ).font.size = Pt(9)

    copy_p = doc.add_paragraph()
    cr = copy_p.add_run(f"© {year} {s.company_name} Inc. All Rights Reserved.")
    cr.font.size = Pt(8.5)
    cr.font.color.rgb = RGBColor(0x6B, 0x6F, 0x63)

    doc.save(str(out_path))
    return out_path
